import pandas as pd
import os
import fitz
from docx import Document
from pathlib import Path
import time
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
import pythoncom
import win32com.client
from collections import defaultdict

# Add at top of Merge.py
import logging

def _validate_pdf(path: str) -> tuple[bool, str]:
    """Returns (is_valid, reason). Cheap check before fitz.open."""
    if not path:
        return False, "empty path"
    if not os.path.exists(path):
        return False, "file does not exist"
    try:
        size = os.path.getsize(path)
    except OSError as e:
        return False, f"stat failed: {e}"
    if size == 0:
        return False, "zero bytes (likely failed conversion)"
    if size < 100:
        return False, f"suspiciously small ({size} bytes)"
    return True, ""

class Preacher:
    def __init__(self,pNum:str,pName:str):
        self.pNum = pNum
        self.pName = pName
        self.note = ""
        self.report = ""
        self.additonal:list[str] = []
        self.noteNecessary = False

        if self.pNum.endswith("c"):
            self.type = "child"
        elif self.pNum.endswith("w"):
            self.type = "widow"
        else:
            self.type = "default"

    def enforce(self):
        self.noteNecessary = True
    
    def ready(self)->bool:
        return ((not self.noteNecessary or self.note!="") 
                and (self.report!="")
                ) 
    def getFiles(self):
        files = [self.report, self.note]
        return [file for file in files if file!=""]         
    
    def display(self):
        print()
        print("***Preacher Printing***")
        print("Name:",self.pName)
        print("Preacher Number:",self.pNum)
        print("Preacher Type:",self.type)
        print("Report:",self.report)
        print("TY note necessary:",self.noteNecessary)
        print("TY note:", self.note)
    def reasons(self):
        reasons = []
        if (self.noteNecessary and self.note==""):
            reasons.append(f"Missing TY note")
        if (self.report == ""):
            reasons.append(f"Missing report")
        return reasons
        
class Donor:
    def __init__(self, pNum = "", pName = "", eNum = "", eName = "", street = "", city = "", state = "", zip = "", email = "",QReport="", aNum = "", *args, **kwargs):
        self.eNum = eNum
        if "\n" in eName:
            self.eName = eName.split("\n")[0]
            self.additional = eName.split("\n")[1]
        else:
            self.eName = eName
            self.additional = ""
        self.street = street
        self.city = city
        self.state = state
        self.zip = zip
        self.email = email
        self.send_quarterly = (True if QReport == "E-Mail" else False)
        self.aNum = aNum
        self.preachers = {pNum: Preacher(pNum,pName)}
        self.coverLetter = ""
        self.extra_notes = []
    
    def __eq__(self, other):
        if not isinstance(other,Donor):
            return NotImplemented
        return self.aNum == other.aNum
    
    def __hash__(self):
        return hash((self.aNum))
    
    def addPreacher(self,pNum,pName):
        if pNum not in self.preachers:
            self.preachers[pNum] = Preacher(pNum, pName)

    def ready(self):
        for preacher in self.preachers.values():
            if not preacher.ready():
                return False
        if self.coverLetter == "":
            return False
        else:
            return True
        
    def getFiles(self) -> list[str]:
        files = []
        files.append(self.coverLetter)
        for preacher in self.preachers.values():
            files.extend(preacher.getFiles())
        
        for note in self.extra_notes:
            files.append(note)
        
        return [file for file in files if file!='']
    
    def display(self):
        print()
        print("========Donor Printing========")
        print("Envelope Number:",self.eNum)
        print("Envelope Name:",self.eName)
        print("Street:",self.street)
        print("City:",self.city)
        print("State:",self.state)
        print("ZIP Code:",self.zip)
        print("Email Address:",self.email)
        print("Account Number:",self.aNum)
        print("Coverletter:", self.coverLetter)
        print("Extra Notes/Reports:", ", ".join(self.extra_notes))
        print("Additional:",self.additional)
        for preacher in self.preachers.values():
            preacher.display()
    
    def reasons(self):
        reasons = {}
        if self.coverLetter=="":
            reasons[f"D#{self.eNum} (A#{self.aNum})"] = "Missing coversheet"
        for preacher in self.preachers.values():
            if not preacher.ready():
                reasons[preacher.pNum] = preacher.reasons()
        return reasons

def getDonors(file:str=".\\Data\\Spreadsheets\\2026-01-22 Cover Sheet Data.xlsx") -> dict[str,Donor]:
    coverSheetDF = pd.read_excel(file)
    dropColumns = ['Type', 'Status', 'Note', 'Sponsorship Amount', 'Pledge','Donor/Preacher Combo']
    coverSheetDF = coverSheetDF.dropna(subset=['Account Number'])
    coverSheetDF["Account Number"] = coverSheetDF["Account Number"].astype(int).astype(str) 
    coverSheetDF = (coverSheetDF.dropna(subset=['Account Number'])
                    .astype(str)
                    .drop(columns=dropColumns)
                    .fillna("")
                    .rename(columns={
                        "Preacher Number": "pNum",
                        "Preacher Name": "pName",
                        "Envelope Number": "eNum",
                        "Envelope Name": "eName",
                        "Primary Street": "street",
                        "Primary City": "city",
                        "Primary State": "state",
                        "Primary ZIP Code": "zip",
                        "Primary Email Address": "email",
                        "Account Number": "aNum",
                        'Send Quarterly Report Via': "QReport",
                    }).astype(str)
                    
                    )
    coverSheetDF["street"] = coverSheetDF["street"].str.replace("\n",' ')
    coverSheetDF["eName"] = coverSheetDF["eName"].str.replace("\n",' ')

    donors = {}
    for _, row in coverSheetDF.iterrows():
        row = {str(k): v for k, v in row.fillna("").to_dict().items()}
        aNum = str(row["aNum"])

        if aNum not in donors:
            donors[aNum] = Donor(**row)
        else:
            donors[aNum].addPreacher(str(row["pNum"]), str(row["pName"]))

    # For every child preacher (pNum ending in "c"), ensure the donor also has
    # a placeholder for the base preacher so addReports can assign its report.
    for donor in donors.values():
        children_nums = [preacher.pNum[:-1] for _, preacher in donor.preachers.items() if preacher.type == "child"]
        for num in children_nums:
            if num not in donor.preachers:
                donor.preachers[num] = Preacher(num, f"{num} Guardian")
        # for _, preacher in donor.preachers.items():
        #     if preacher.type == "child":
        #         base_pNum = preacher.pNum[:-1]
        #         if base_pNum not in donor.preachers:
        #             donor.preachers[base_pNum] = Preacher(base_pNum, f"{preacher.pName} Guardian")

    return donors

def findDonor(*, eNum=None, eName=None, aNum=None, donors=None) -> Donor | None:
    if donors is None:
        return None
    for donor in donors.values():
        if (
            (eNum  and donor.eNum  == eNum) or
            (eName and donor.eName == eName) or
            (aNum  and donor.aNum  == aNum)
        ):
            return donor

    return None


def enforceTY(donors:dict[str,Donor],extraGiftFile:str = "") -> None:
    # Get data from both sheets, rename columns, and concatenate into one DataFrame
    giftDF = pd.read_excel(extraGiftFile,sheet_name="Spons Xtra")
    giftDF = giftDF.rename(columns={  # ID	Name	Amount	Donor #	Donor Name
        "ID": "pNum",
        "Donor #": "eNum"
    })[["eNum","pNum"]].dropna().astype(str)
    
    
    donors_by_eNum = {d.eNum: d for d in donors.values() if d.eNum}
    sponsorBlacklist = ["21956","20325","1940","1287","1565","1787","3548","3628","3990","5705","7608","604110"]
    
    for _, row in giftDF.iterrows():
        if row["eNum"] in sponsorBlacklist:
             continue
        donor = donors_by_eNum.get(row["eNum"])
        if donor == None:
                print("what the.")
                print(str(row))
                continue #fix this
        preacher = donor.preachers.get(row["pNum"])
        if preacher is not None:
            preacher.enforce()
        
def addReports(donors:dict[str,Donor],directory:str=".\\Data\\Reports") -> None:
    reports = os.listdir(directory)
    blacklist=["pcf","widow","rdf",".docx","rd"]

    donors_by_eNum = {d.eNum: d for d in donors.values() if d.eNum}
    donors_by_pNum = defaultdict(list)
    for d in donors.values():
        for pNum in d.preachers:
            donors_by_pNum[pNum].append(d)

    for report in reports:

        if any(kw in report for kw in blacklist):
            continue
        
        path = os.path.join(directory, report)

        if "ty" in report:
            pNum, eNum = report[:-4].replace(' ','').split('ty')
            donor = donors_by_eNum.get(eNum)
            if donor is None:
                print("yuck donor,",report) #donor referenced was not found at all                
            else:
                tempPreacher = donor.preachers.get(pNum)
                if tempPreacher is not None:
                    tempPreacher.note = path
                else:
                    donor.extra_notes.append(path)
                    #preacher is not listed on donors donations, still added though
        else:

            pNum = report[:-4]
            for donor in donors_by_pNum.get(pNum, []):
                if donor.preachers[pNum] is not None:
                    donor.preachers[pNum].report = path

def merge(files, output_path: str):
    if not files:
        raise ValueError("merge() called with no files")

    # Validate every input before we start
    bad = []
    for f in files:
        ok, reason = _validate_pdf(f)
        if not ok:
            bad.append(f"  {f}: {reason}")
    if bad:
        raise FileNotFoundError(
            "Cannot merge — invalid input files:\n" + "\n".join(bad)
        )

    # Make sure output dir exists
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

    merged_pdf = fitz.open()
    try:
        for file in files:
            pdf = fitz.open(file)
            merged_pdf.insert_pdf(pdf)
            pdf.close()
        merged_pdf.save(output_path)
    finally:
        merged_pdf.close()

    # Verify it actually wrote
    ok, reason = _validate_pdf(output_path)
    if not ok:
        raise IOError(f"Merge appeared to succeed but output is invalid: {reason}")
    
def mergeDonors(donors:list[Donor],output_path):

    for donor in donors:
        files = list(dict.fromkeys(donor.getFiles()))
        path = os.path.join(output_path, f"{'d'+donor.eNum if donor.eNum else 'a'+donor.aNum}{'e' if donor.send_quarterly else ''}.pdf")
        merge(files, path)


def fill_template(input_path, output_path, replacements:dict):
    """
    Replace placeholder text in a DOCX file while preserving formatting.

    Placeholders may be split across several runs (e.g. «Envelope_Name» is
    stored as the runs « | Envelope_Name | »). For this template they only
    ever appear in plain body paragraphs or in table cells, so those are the
    only locations scanned.

    The replacement is done in place at the <w:t> (text node) level: runs are
    never removed or rebuilt, so run formatting and non-text run content such
    as tabs (<w:tab/>), line breaks and drawings are left untouched. A value
    containing '\\n' is rendered as line breaks within its run.

    Args:
        input_path (str): Path to the input DOCX file.
        output_path (str): Path where the modified DOCX will be saved.
        replacements (dict): Mapping of placeholder text -> replacement text,
                             e.g. {'«Envelope_Name»': 'John Doe'}.
    """
    doc = Document(input_path)

    def text_nodes(para_element):
        """The <w:t> nodes of a paragraph's direct runs, in document order.

        Only direct-child runs are inspected so that nested content (e.g. a
        text box anchored inside the paragraph) is never modified.
        """
        return [t
                for run in para_element.findall(qn('w:r'))
                for t in run.findall(qn('w:t'))]

    def write_value(t_elem, value):
        """Set a <w:t>'s text to value, expanding '\\n' into <w:br> breaks.

        Extra lines are inserted as siblings inside the same run so they keep
        the run's formatting.
        """
        t_elem.set(qn('xml:space'), 'preserve')
        if '\n' not in value:
            t_elem.text = value
            return

        run = t_elem.getparent()
        insert_at = list(run).index(t_elem)
        lines = value.split('\n')
        t_elem.text = lines[0]
        for offset, line in enumerate(lines[1:], start=1):
            br = OxmlElement('w:br')
            nt = OxmlElement('w:t')
            nt.set(qn('xml:space'), 'preserve')
            nt.text = line
            run.insert(insert_at + (offset * 2) - 1, br)
            run.insert(insert_at + (offset * 2), nt)

    def replace_in_paragraph(para_element):
        """Replace every placeholder occurrence in a single paragraph."""
        # One replacement per pass; re-scan afterwards so placeholders that
        # span runs (and repeated placeholders) are all resolved.
        while True:
            nodes = text_nodes(para_element)
            if not nodes:
                return
            texts = [t.text or '' for t in nodes]
            full = ''.join(texts)

            # Earliest-starting placeholder present in this paragraph.
            match = None
            for key in replacements:
                pos = full.find(key)
                if pos != -1 and (match is None or pos < match[0]):
                    match = (pos, key)
            if match is None:
                return

            start, key = match
            end = start + len(key)
            value = replacements[key]

            # Character span [node_start, node_end) covered by each <w:t>.
            bounds = []
            cursor = 0
            for txt in texts:
                bounds.append((cursor, cursor + len(txt)))
                cursor += len(txt)

            # Nodes that the placeholder overlaps.
            covered = [i for i, (s, e) in enumerate(bounds)
                       if s < end and e > start]
            first, last = covered[0], covered[-1]

            prefix = texts[first][:start - bounds[first][0]]
            suffix = texts[last][end - bounds[last][0]:]

            if first == last:
                write_value(nodes[first], prefix + value + suffix)
            else:
                write_value(nodes[first], prefix + value)
                for i in covered[1:-1]:
                    nodes[i].text = ''
                nodes[last].text = suffix
                nodes[last].set(qn('xml:space'), 'preserve')

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph._element)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph._element)

    doc.save(output_path)


def createCoverLetter(donor:Donor,output_path=".\\Data\\Output\\CoverLetters",template_path=".\\Data\\Spreadsheets\\Cover Sheet Q3-2025.docx"):
    pdf_target  = os.path.join(output_path, f"{donor.aNum}.pdf")
    docx_target = os.path.join(output_path, f"{donor.aNum}.docx")

    # Only skip if a real PDF already exists
    ok, _ = _validate_pdf(pdf_target)
    if ok:
        print(f"{donor.aNum} CL Skipped (PDF exists)")
        return

    # Stale orphan docx from a prior failed run? Nuke it.
    if os.path.exists(docx_target):
        try:
            os.remove(docx_target)
        except OSError:
            pass
    if f"{donor.aNum}.docx" in os.listdir(output_path):
        print(f"{donor.aNum} CL Skipped!")
        return 
    
    data = {
        "«Envelope_Name»": donor.eName,
        "«Additional»": f"{'\n' if donor.additional!= '' else ''}{donor.additional}",
        "«Primary_Street»": donor.street,
        "«Primary_City»": donor.city,
        "«Primary_State»": donor.state,
        "«Primary_ZIP_Code»": donor.zip,
        "«Envelope_Number»": donor.eNum,
    }
    for i in range(1,14):
        if i < len(donor.preachers)+1:
            preacher = list(donor.preachers.values())[i-1]
            data[f"«Preacher_{i}»"] = preacher.pNum+' - '+preacher.pName
        else:
            data[f"«Preacher_{i}»"] = ""
    path = os.path.join(output_path, f"{donor.aNum}.docx")
    
    fill_template(template_path,path,data)
    print(f"{donor.aNum} created!")

def docx_to_pdf(input_dir, output_dir, restart_every=50, pause=0.5, files=None):
    input_dir = Path(input_dir).resolve()
    output_dir = Path(output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    if files is None:
        files = sorted(input_dir.glob("*.docx"))
    else:
        files = [Path(f) for f in files]

    # Skip files whose PDFs already exist before touching Word
    pending = [f for f in files if not (output_dir / f"{f.stem}.pdf").exists()]
    if not pending:
        return
    
    wdFormatPDF = 17
    
    pythoncom.CoInitialize()
    word = None

    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

        count = 0

        for idx, docx_path in enumerate(files, 1):
            pdf_path = output_dir / f"{docx_path.stem}.pdf"

            if pdf_path.exists():
                continue

            doc = None
            try:
                print(f"[{idx}/{len(files)}] {docx_path.name}")
                doc = word.Documents.Open(str(docx_path))
                doc.SaveAs(str(pdf_path), FileFormat=wdFormatPDF)
                count += 1
            except Exception as e:
                print(f"Failed: {docx_path.name}: {e}")
            finally:
                if doc is not None:
                    doc.Close(False)

            if count >= restart_every:
                word.Quit()
                time.sleep(pause)
                word = win32com.client.DispatchEx("Word.Application")
                word.Visible = False
                word.DisplayAlerts = 0
                count = 0

    finally:
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()
def assignCoverLetters(donors, directory=".\\Data\\Output\\CoverLetters"):
    if not os.path.isdir(directory):
        print(f"[assignCoverLetters] directory missing: {directory}")
        return
    assigned = skipped = 0
    for fname in os.listdir(directory):
        if not fname.lower().endswith(".pdf"):
            continue
        full = os.path.join(directory, fname)
        ok, reason = _validate_pdf(full)
        if not ok:
            print(f"[assignCoverLetters] skipping {fname}: {reason}")
            skipped += 1
            continue
        aNum = fname[:-4]
        if aNum in donors:
            donors[aNum].coverLetter = full
            assigned += 1
    print(f"[assignCoverLetters] assigned={assigned} skipped={skipped}")
            
if __name__ == "__main__":
    output_dir = ".\\Data\\Output" 
    cover_letter_dir = output_dir+"\\CoverLetters"
    coverSheetFile = ".\\Data\\Spreadsheets\\2026-01-22 Cover Sheet Data.xlsx"
    extraGiftFile = ".\\Data\\Spreadsheets\\2026-01-22 Qtr 3 2025 extra gift reports.xlsx"
    donors = getDonors(coverSheetFile)
     
    # for donor in donors.values():
    #     createCoverLetter(donor,cover_letter_dir)
    # docx_to_pdf(".\\Data\\Output\\CoverLetters",".\\Data\\Output\\CoverLetters")
    assignCoverLetters(donors,cover_letter_dir)

    enforceTY(donors,extraGiftFile)


    addReports(donors)

    readyDonors = [donor for donor in donors.values() if donor.ready()]
    
    mergeDonors(readyDonors,".\\Data\\Output\\Reports")

